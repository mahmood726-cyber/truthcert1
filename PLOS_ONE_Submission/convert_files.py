"""
PLOS ONE Submission File Converter
Converts markdown files to DOCX and PDF, creates ZIP archive, and placeholder figures
"""

import os
import re
import zipfile
from pathlib import Path

# Document conversion libraries
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from fpdf import FPDF
from PIL import Image, ImageDraw, ImageFont

# Base directory
BASE_DIR = Path(r"C:\Truthcert1\PLOS_ONE_Submission")

def clean_markdown(text):
    """Remove markdown formatting symbols"""
    # Remove headers (# symbols)
    text = re.sub(r'^#{1,6}\s*', '', text, flags=re.MULTILINE)
    # Remove bold (**text** or __text__)
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'__([^_]+)__', r'\1', text)
    # Remove italic (*text* or _text_)
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    text = re.sub(r'_([^_]+)_', r'\1', text)
    # Remove code blocks
    text = re.sub(r'```[^`]*```', '', text, flags=re.DOTALL)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    # Remove horizontal rules
    text = re.sub(r'^---+$', '', text, flags=re.MULTILINE)
    text = re.sub(r'^\*\*\*+$', '', text, flags=re.MULTILINE)
    return text

def parse_markdown_table(table_text):
    """Parse a markdown table into rows and columns"""
    lines = [l.strip() for l in table_text.strip().split('\n') if l.strip()]
    rows = []
    for line in lines:
        if re.match(r'^[\|\-\s:]+$', line):  # Skip separator lines
            continue
        cells = [c.strip() for c in line.split('|')]
        cells = [c for c in cells if c]  # Remove empty cells from edges
        if cells:
            rows.append(cells)
    return rows

def add_table_to_doc(doc, table_data):
    """Add a formatted table to a Word document"""
    if not table_data or len(table_data) < 1:
        return

    num_cols = max(len(row) for row in table_data)
    table = doc.add_table(rows=len(table_data), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(table_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = row.cells[j]
                cell.text = clean_markdown(cell_text)
                # Bold header row
                if i == 0:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True

def convert_md_to_docx(md_path, docx_path):
    """Convert markdown file to Word document with proper formatting"""
    print(f"Converting {md_path.name} to DOCX...")

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Process content
    lines = content.split('\n')
    i = 0
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i]

        # Check for table start
        if '|' in line and not in_table:
            in_table = True
            table_lines = [line]
            i += 1
            continue

        # Continue table
        if in_table:
            if '|' in line:
                table_lines.append(line)
                i += 1
                continue
            else:
                # End of table
                table_text = '\n'.join(table_lines)
                table_data = parse_markdown_table(table_text)
                add_table_to_doc(doc, table_data)
                doc.add_paragraph()  # Add space after table
                in_table = False
                table_lines = []

        # Skip empty lines (but add paragraph break)
        if not line.strip():
            i += 1
            continue

        # Skip code blocks
        if line.strip().startswith('```'):
            while i < len(lines) and not lines[i].strip().endswith('```'):
                i += 1
            i += 1
            continue

        # Handle headers
        header_match = re.match(r'^(#{1,6})\s*(.+)$', line)
        if header_match:
            level = len(header_match.group(1))
            text = clean_markdown(header_match.group(2))

            if level == 1:
                para = doc.add_heading(text, level=1)
            elif level == 2:
                para = doc.add_heading(text, level=2)
            elif level == 3:
                para = doc.add_heading(text, level=3)
            else:
                para = doc.add_paragraph()
                run = para.add_run(text)
                run.bold = True
            i += 1
            continue

        # Handle horizontal rules
        if re.match(r'^---+$', line) or re.match(r'^\*\*\*+$', line):
            i += 1
            continue

        # Regular paragraph
        clean_line = clean_markdown(line)
        if clean_line.strip():
            para = doc.add_paragraph(clean_line)

        i += 1

    # Handle any remaining table
    if in_table and table_lines:
        table_text = '\n'.join(table_lines)
        table_data = parse_markdown_table(table_text)
        add_table_to_doc(doc, table_data)

    doc.save(docx_path)
    print(f"  Created: {docx_path.name}")

def convert_txt_to_docx(txt_path, docx_path):
    """Convert plain text file to Word document"""
    print(f"Converting {txt_path.name} to DOCX...")

    with open(txt_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Add content paragraph by paragraph
    for para_text in content.split('\n\n'):
        if para_text.strip():
            doc.add_paragraph(para_text.strip())

    doc.save(docx_path)
    print(f"  Created: {docx_path.name}")

class PDF(FPDF):
    def header(self):
        self.set_font('Helvetica', 'B', 10)
        self.cell(0, 10, 'TruthCert-PairwisePro - Supporting Information', align='C', new_x='LMARGIN', new_y='NEXT')
        self.ln(5)

    def footer(self):
        self.set_y(-15)
        self.set_font('Helvetica', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}', align='C')

def convert_md_to_pdf(md_path, pdf_path):
    """Convert markdown file to PDF"""
    print(f"Converting {md_path.name} to PDF...")

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    pdf = PDF()
    pdf.set_margins(10, 25, 10)  # Left, top, right margins
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Process content
    lines = content.split('\n')

    for line in lines:
        # Skip code block markers
        if line.strip().startswith('```'):
            continue

        # Handle headers
        header_match = re.match(r'^(#{1,6})\s*(.+)$', line)
        if header_match:
            level = len(header_match.group(1))
            text = clean_markdown(header_match.group(2))

            if level == 1:
                pdf.set_font('Helvetica', 'B', 16)
            elif level == 2:
                pdf.set_font('Helvetica', 'B', 14)
            elif level == 3:
                pdf.set_font('Helvetica', 'B', 12)
            else:
                pdf.set_font('Helvetica', 'B', 11)

            pdf.ln(5)
            pdf.multi_cell(0, 8, text)
            pdf.ln(3)
            continue

        # Skip horizontal rules
        if re.match(r'^---+$', line) or re.match(r'^\*\*\*+$', line):
            pdf.ln(5)
            continue

        # Handle tables (simplified - just show as text)
        if '|' in line:
            pdf.set_font('Courier', '', 6)
            # Clean and truncate line to fit page width
            clean_line = line.strip()
            # Remove pipe characters and format nicely
            cells = [c.strip() for c in clean_line.split('|') if c.strip()]
            clean_line = '  '.join(cells[:6])  # Max 6 columns
            if len(clean_line) > 90:
                clean_line = clean_line[:90] + '...'
            if clean_line and not re.match(r'^[\-\s:]+$', clean_line):
                try:
                    pdf.cell(0, 4, clean_line, new_x='LMARGIN', new_y='NEXT')
                except:
                    pass  # Skip problematic lines
            continue

        # Regular text
        clean_line = clean_markdown(line)
        if clean_line.strip():
            pdf.set_font('Helvetica', '', 10)
            # Split long lines into chunks
            words = clean_line.split()
            current_line = ""
            for word in words:
                test_line = current_line + " " + word if current_line else word
                if len(test_line) > 95:
                    if current_line:
                        try:
                            pdf.cell(0, 5, current_line, new_x='LMARGIN', new_y='NEXT')
                        except:
                            pass
                    current_line = word
                else:
                    current_line = test_line
            if current_line:
                try:
                    pdf.cell(0, 5, current_line, new_x='LMARGIN', new_y='NEXT')
                except:
                    pass
        else:
            pdf.ln(3)

    pdf.output(pdf_path)
    print(f"  Created: {pdf_path.name}")

def create_zip_archive(files, zip_path):
    """Create ZIP archive from list of files"""
    print(f"Creating ZIP archive...")

    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for file_path in files:
            if file_path.exists():
                zipf.write(file_path, file_path.name)
                print(f"  Added: {file_path.name}")

    print(f"  Created: {zip_path.name}")

def create_placeholder_figure(fig_num, title, width_in, height_in, output_path):
    """Create a placeholder TIFF figure at 300 DPI"""
    print(f"Creating Figure {fig_num}...")

    # Calculate pixel dimensions at 300 DPI
    dpi = 300
    width_px = int(width_in * dpi)
    height_px = int(height_in * dpi)

    # Create image with white background
    img = Image.new('RGB', (width_px, height_px), color='white')
    draw = ImageDraw.Draw(img)

    # Draw border
    draw.rectangle([10, 10, width_px-10, height_px-10], outline='black', width=3)

    # Add text
    try:
        font_large = ImageFont.truetype("arial.ttf", 60)
        font_small = ImageFont.truetype("arial.ttf", 36)
    except:
        font_large = ImageFont.load_default()
        font_small = ImageFont.load_default()

    # Title
    title_text = f"Figure {fig_num}"
    draw.text((width_px//2, height_px//4), title_text, fill='black', font=font_large, anchor='mm')

    # Description
    draw.text((width_px//2, height_px//2), title, fill='gray', font=font_small, anchor='mm')

    # Dimensions note
    dim_text = f"{width_in}\" x {height_in}\" at 300 DPI"
    draw.text((width_px//2, height_px*3//4), dim_text, fill='gray', font=font_small, anchor='mm')

    # Placeholder note
    note_text = "[Replace with actual screenshot/diagram]"
    draw.text((width_px//2, height_px*7//8), note_text, fill='red', font=font_small, anchor='mm')

    # Save as TIFF
    img.save(output_path, format='TIFF', dpi=(300, 300), compression='tiff_lzw')
    print(f"  Created: {output_path.name} ({os.path.getsize(output_path) / 1024:.1f} KB)")

def main():
    print("=" * 60)
    print("PLOS ONE Submission File Converter")
    print("=" * 60)
    print()

    # 1. Convert Manuscript to DOCX
    print("1. Converting Manuscript to DOCX...")
    manuscript_md = BASE_DIR / "03_Manuscript_TruthCert_PairwisePro.md"
    manuscript_docx = BASE_DIR / "03_Manuscript_TruthCert_PairwisePro.docx"
    if manuscript_md.exists():
        convert_md_to_docx(manuscript_md, manuscript_docx)
    else:
        print(f"  ERROR: {manuscript_md} not found")
    print()

    # 2. Convert Cover Letter to DOCX
    print("2. Converting Cover Letter to DOCX...")
    cover_txt = BASE_DIR / "01_Cover_Letter.txt"
    cover_docx = BASE_DIR / "01_Cover_Letter.docx"
    if cover_txt.exists():
        convert_txt_to_docx(cover_txt, cover_docx)
    else:
        print(f"  ERROR: {cover_txt} not found")
    print()

    # 3. Convert Supporting Files to PDF
    print("3. Converting Supporting Information to PDF...")
    supporting_files = [
        "S1_Technical_Documentation.md",
        "S2_Monte_Carlo_Protocol.md",
        "S4_HTA_Validation.md",
        "S5_User_Guide.md"
    ]

    for sf in supporting_files:
        md_path = BASE_DIR / sf
        pdf_path = BASE_DIR / sf.replace('.md', '.pdf')
        if md_path.exists():
            convert_md_to_pdf(md_path, pdf_path)
        else:
            print(f"  ERROR: {md_path} not found")
    print()

    # 4. Create ZIP archive for S3 files
    print("4. Creating S3 ZIP archive...")
    s3_files = [
        BASE_DIR / "S3_R_Validation_Code.R",
        BASE_DIR / "S3_R_Validation_Output.txt"
    ]
    s3_zip = BASE_DIR / "S3_R_Validation.zip"
    create_zip_archive(s3_files, s3_zip)
    print()

    # 5. Create placeholder figures
    print("5. Creating placeholder TIFF figures...")
    figures = [
        (1, "User Interface Overview", 7.5, 6.0),
        (2, "Verdict Classification Flowchart", 7.5, 8.0),
        (3, "Validation Results", 7.5, 6.0),
        (4, "HTA Module Interface", 5.2, 5.0),
        (5, "Complete Workflow", 7.5, 3.5)
    ]

    for fig_num, title, width, height in figures:
        output_path = BASE_DIR / f"Figure_{fig_num}.tiff"
        create_placeholder_figure(fig_num, title, width, height, output_path)
    print()

    # Summary
    print("=" * 60)
    print("CONVERSION COMPLETE")
    print("=" * 60)
    print()
    print("Files created:")

    expected_files = [
        "03_Manuscript_TruthCert_PairwisePro.docx",
        "01_Cover_Letter.docx",
        "S1_Technical_Documentation.pdf",
        "S2_Monte_Carlo_Protocol.pdf",
        "S4_HTA_Validation.pdf",
        "S5_User_Guide.pdf",
        "S3_R_Validation.zip",
        "Figure_1.tiff",
        "Figure_2.tiff",
        "Figure_3.tiff",
        "Figure_4.tiff",
        "Figure_5.tiff"
    ]

    for fname in expected_files:
        fpath = BASE_DIR / fname
        if fpath.exists():
            size = os.path.getsize(fpath) / 1024
            print(f"  [OK] {fname} ({size:.1f} KB)")
        else:
            print(f"  [MISSING] {fname}")

    print()
    print("Next steps:")
    print("1. Replace placeholder figures with actual screenshots/diagrams")
    print("2. Review DOCX formatting and adjust as needed")
    print("3. Upload to PLOS ONE Editorial Manager")

if __name__ == "__main__":
    main()
